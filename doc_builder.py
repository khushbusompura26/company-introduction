"""
doc_builder.py — Formatted Hinglish video script using python-docx
Cover page → How to use → Slide sections → Production notes
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY = RGBColor(0x1B, 0x3A, 0x6B)
GOLD = RGBColor(0xC8, 0xA9, 0x51)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x77, 0x77, 0x77)


def _para(doc, text='', size=12, color=None, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
          font_name='Calibri'):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size      = Pt(size)
        r.font.bold      = bold
        r.font.italic    = italic
        r.font.name      = font_name
        r.font.color.rgb = color or DARK
    return p


def _heading(doc, text, level=1):
    """Styled heading without default Word heading styles"""
    size  = 20 if level == 1 else 14
    color = NAVY if level == 1 else GOLD
    fname = 'Georgia' if level == 1 else 'Calibri'
    return _para(doc, text, size=size, color=color, bold=True,
                 space_before=12, space_after=6, font_name=fname)


def build_script_document(data: dict) -> bytes:
    """
    Build a fully formatted DOCX video script.
    Returns raw DOCX bytes.
    """
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width      = Inches(8.5)
    sec.page_height     = Inches(11)
    sec.left_margin     = Inches(1.0)
    sec.right_margin    = Inches(1.0)
    sec.top_margin      = Inches(1.0)
    sec.bottom_margin   = Inches(1.0)

    cn     = data.get('company_name', 'Company')
    wu     = data.get('website', '')
    slides = data.get('slides', [])
    notes  = data.get('production_notes', {})

    # ── Cover page ─────────────────────────────────────────────
    _para(doc, cn, size=32, color=NAVY, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=60, space_after=12, font_name='Georgia')

    _para(doc, 'Company Introduction \u2014 Video Script',
          size=18, color=RGBColor(0x33, 0x33, 0x33),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    _para(doc, 'Slide-by-Slide Narration for Corporate Video',
          size=13, color=GREY, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    runtime = notes.get('total_runtime', '20-25 minutes')
    _para(doc, f'Total Slides: {len(slides)}  \u00b7  Estimated Runtime: {runtime}',
          size=12, color=RGBColor(0x55, 0x55, 0x55),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    _para(doc, wu, size=12, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()

    # ── How to use ─────────────────────────────────────────────
    _heading(doc, 'HOW TO USE THIS SCRIPT')
    _para(doc,
          'This document is divided into sections \u2014 one for each slide. '
          'Each section contains:\n'
          '\u2022 ON SCREEN \u2014 the key pointers/visuals visible on that slide\n'
          '\u2022 VIDEO SCRIPT \u2014 the exact narration (in Hinglish)\n'
          '\u2022 ESTIMATED SPEAKING TIME \u2014 approximate duration\n\n'
          'Read the ON SCREEN pointers to understand the slide, '
          'then speak the VIDEO SCRIPT narration naturally and conversationally.',
          size=12, color=DARK, space_after=6)

    doc.add_page_break()

    # ── Slide sections ─────────────────────────────────────────
    for sl in slides:
        num   = sl.get('slide_number', '?')
        title = sl.get('title', 'Slide')
        pts   = sl.get('key_points', [])
        script = sl.get('script', '[No script generated for this slide]')

        # Slide header
        _heading(doc, f'SLIDE {num} \u2014 {title}')

        # ON SCREEN
        _para(doc, '\U0001f5a5  ON SCREEN (Slide Pointers)',
              size=13, color=GOLD, bold=True, space_before=4, space_after=4)

        for pt in pts:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run('\u25b8  ' + str(pt))
            r.font.size      = Pt(12)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r.font.name      = 'Calibri'

        # VIDEO SCRIPT
        _para(doc, '\U0001f3a4  VIDEO SCRIPT \u2014 Narration (Speak These Words)',
              size=13, color=NAVY, bold=True, space_before=10, space_after=4)

        _para(doc, script, size=12, color=DARK, space_after=4)

        # Speaking time
        _para(doc, '\u23f1  Estimated speaking time: \u223c50 seconds (1 min)',
              size=10, color=GREY, italic=True, space_after=16)

    # ── Production notes ───────────────────────────────────────
    doc.add_page_break()
    _heading(doc, 'PRODUCTION NOTES')

    note_items = [
        ('Total Estimated Video Length', notes.get('total_runtime', '')),
        ('Recommended Speaking Pace',    notes.get('pace', '')),
        ('Background Music',             notes.get('music_suggestion', '')),
        ('Recording Tips',               notes.get('recording_tips', '')),
    ]
    for label, value in note_items:
        if value:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            rl = p.add_run(label + ': ')
            rl.bold = True
            rl.font.size      = Pt(12)
            rl.font.color.rgb = DARK
            rl.font.name      = 'Calibri'
            rv = p.add_run(value)
            rv.font.size      = Pt(12)
            rv.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            rv.font.name      = 'Calibri'

    # ── Save ───────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
